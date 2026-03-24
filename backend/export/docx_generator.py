"""
DOCX Generator with IEEE Formatting
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os
import re

from config import Config


class IEEEDOCXGenerator:
    """Generate IEEE-formatted DOCX documents"""
    
    def __init__(self, output_dir: str = None):
        self.output_dir = output_dir or Config.OUTPUT_DIR
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate(self, research_data: dict, filename: str = None) -> str:
        """
        Generate DOCX from research data
        
        Args:
            research_data: Dict with topic, synthesis, bibliography, metadata
            filename: Output filename (auto-generated if None)
        
        Returns:
            Path to generated DOCX
        """
        if filename is None:
            topic_slug = research_data['topic'].replace(' ', '_')[:50]
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"research_{topic_slug}_{timestamp}.docx"
        
        filepath = os.path.join(self.output_dir, filename)
        
        # Create document
        doc = Document()
        
        # Setup styles
        self._setup_styles(doc)
        
        # Setup page
        self._setup_page(doc)
        
        # Title
        title = doc.add_paragraph(research_data['topic'])
        title.style = 'Title'
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Author and date
        author = doc.add_paragraph("Deep Research Assistant")
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author.runs[0].font.size = Pt(12)
        
        date = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date.runs[0].font.size = Pt(11)
        
        doc.add_paragraph()  # Spacing
        
        # Parse and add content
        synthesis = research_data.get('synthesis', '')
        sections = self._parse_sections(synthesis)
        
        for i, (section_title, section_content) in enumerate(sections):
            # Section heading
            if section_title.upper() != 'ABSTRACT':
                section_num = self._get_section_number(i)
                heading_text = f"{section_num}. {section_title.upper()}"
            else:
                heading_text = section_title.upper()
            
            heading = doc.add_heading(heading_text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Section content
            paragraphs = section_content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    # Clean up paragraph text
                    cleaned_text = para_text.strip()
                    # Remove markdown bold formatting (**text**)
                    cleaned_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cleaned_text)
                    # Remove any stray asterisks
                    cleaned_text = cleaned_text.replace('**', '')
                    
                    para = doc.add_paragraph(cleaned_text)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.paragraph_format.space_after = Pt(6)
                    
                    # Set font
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
        
        # Bibliography
        bibliography = research_data.get('bibliography', '')
        if bibliography:
            doc.add_paragraph()
            
            ref_heading = doc.add_heading('REFERENCES', level=1)
            ref_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Parse references
            ref_lines = bibliography.split('\n')
            for line in ref_lines:
                if line.strip() and line.strip() != 'REFERENCES':
                    ref_para = doc.add_paragraph(line.strip())
                    ref_para.paragraph_format.left_indent = Inches(0.25)
                    ref_para.paragraph_format.first_line_indent = Inches(-0.25)
                    ref_para.paragraph_format.space_after = Pt(3)
                    
                    for run in ref_para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(9)
        
        # Save document
        doc.save(filepath)
        
        return filepath
    
    def _setup_styles(self, doc):
        """Setup custom styles for IEEE format"""
        styles = doc.styles
        
        # Modify Normal style
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(10)
        
        # Modify Heading 1 style
        heading1_style = styles['Heading 1']
        heading1_font = heading1_style.font
        heading1_font.name = 'Times New Roman'
        heading1_font.size = Pt(10)
        heading1_font.bold = True
        heading1_font.color.rgb = RGBColor(0, 0, 0)
        
        # Modify Title style
        title_style = styles['Title']
        title_font = title_style.font
        title_font.name = 'Times New Roman'
        title_font.size = Pt(24)
        title_font.bold = True
    
    def _setup_page(self, doc):
        """Setup page margins and layout"""
        sections = doc.sections
        for section in sections:
            # Set margins (IEEE format)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(0.625)
            section.right_margin = Inches(0.625)
            
            # Note: Two-column layout requires more complex setup
            # For simplicity, using single column in DOCX
            # Users can manually convert to two-column in Word if needed
    
    def _parse_sections(self, text: str) -> list:
        """Parse text into sections"""
        sections = []
        current_section = None
        current_content = []
        
        lines = text.split('\n')
        
        for line in lines:
            if self._is_section_heading(line):
                if current_section:
                    sections.append((current_section, '\n\n'.join(current_content)))
                
                current_section = line.strip().replace('#', '').strip()
                current_content = []
            else:
                if line.strip():
                    current_content.append(line.strip())
        
        if current_section:
            sections.append((current_section, '\n\n'.join(current_content)))
        
        return sections
    
    def _is_section_heading(self, line: str) -> bool:
        """Check if line is a section heading"""
        line = line.strip()
        
        if line.startswith('#'):
            return True
        
        common_sections = [
            'abstract', 'introduction', 'background', 'methodology',
            'results', 'discussion', 'conclusion', 'references',
            'literature review', 'main findings', 'future work'
        ]
        
        return any(section in line.lower() for section in common_sections)
    
    def _get_section_number(self, index: int) -> str:
        """Get section number in Roman numerals"""
        roman_numerals = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X']
        if index < len(roman_numerals):
            return roman_numerals[index]
        return str(index + 1)
